import streamlit as st
from docx import Document
import difflib
import zipfile
import xml.etree.ElementTree as ET

st.set_page_config(page_title="Compare_Docs", page_icon="📝")

def get_text_from_content_controls(element):
    """
    Extracts text from all content controls (w:sdt elements) inside the given XML element.
    Uses fully qualified tag names instead of xpath with namespaces.
    """
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    texts = []
    for sdt in element.iter("{%s}sdt" % ns):
        sdt_content = sdt.find("{%s}sdtContent" % ns)
        if sdt_content is not None:
            content_text = "".join(sdt_content.itertext()).strip()
            if content_text:
                texts.append(content_text)
    return " ".join(texts)

def get_full_paragraph_text(para):
    """
    Returns the full text of a paragraph, including any text inside content controls.
    """
    full_text = para.text
    try:
        sdt_text = get_text_from_content_controls(para._p)
        if sdt_text:
            full_text += ' ' + sdt_text
    except Exception as e:
        st.error("Error extracting content control text: " + str(e))
    return full_text.strip()

def extract_docx_paragraphs(file):
    """
    Extracts paragraphs as a list of text values.
    """
    document = Document(file)
    return [get_full_paragraph_text(para) for para in document.paragraphs if get_full_paragraph_text(para)]

def extract_docx_tables(file):
    """
    Extracts tables from the document.
    Returns a list of tables where each table is represented as a list of rows,
    and each row is a list of cell texts (including text from content controls).
    """
    document = Document(file)
    tables_data = []
    for table in document.tables:
        table_text = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                try:
                    cell_sdt_text = get_text_from_content_controls(cell._tc)
                    if cell_sdt_text:
                        cell_text += ' ' + cell_sdt_text
                except Exception as e:
                    st.error("Error extracting content control from table cell: " + str(e))
                row_text.append(cell_text.strip())
            table_text.append(row_text)
        tables_data.append(table_text)
    return tables_data

def compare_paragraphs(paragraphs1, paragraphs2):
    """
    Compares two lists of paragraphs and returns a list of diff strings.
    Uses difflib.ndiff to clearly show added or missing text and also displays the full changed value.
    """
    result = []
    max_len = max(len(paragraphs1), len(paragraphs2))
    for i in range(max_len):
        text1 = paragraphs1[i] if i < len(paragraphs1) else ""
        text2 = paragraphs2[i] if i < len(paragraphs2) else ""
        if text1 != text2:
            diff = "\n".join(difflib.ndiff([text1], [text2]))
            full_change = f"Full Change:\n- Original: {text1}\n- Changed: {text2}"
            result.append(f"Paragraph {i+1} text difference:\n{diff}\n{full_change}")
    return result

def compare_tables(tables1, tables2):
    """
    Compares tables from two documents in a cell-by-cell manner and returns detailed diff strings.
    For each cell where a difference is detected, the exact change is displayed along with the full new and old values.
    """
    result = []
    max_table_count = max(len(tables1), len(tables2))
    for t in range(max_table_count):
        table1 = tables1[t] if t < len(tables1) else []
        table2 = tables2[t] if t < len(tables2) else []
        max_rows = max(len(table1), len(table2))
        for row_index in range(max_rows):
            row1 = table1[row_index] if row_index < len(table1) else []
            row2 = table2[row_index] if row_index < len(table2) else []
            max_cells = max(len(row1), len(row2))
            for cell_index in range(max_cells):
                val1 = row1[cell_index] if cell_index < len(row1) else ""
                val2 = row2[cell_index] if cell_index < len(row2) else ""
                if val1 != val2:
                    diff = "\n".join(difflib.ndiff([val1], [val2]))
                    full_change = f"Full Change:\n- Original: {val1}\n- Changed: {val2}"
                    result.append(f"Table {t+1}, Row {row_index+1}, Cell {cell_index+1} difference:\n{diff}\n{full_change}")
            if len(row1) != len(row2):
                result.append(f"Table {t+1}, Row {row_index+1} cell count differs: Original Doc has {len(row1)} cells vs Other Doc2 has {len(row2)} cells.")
        if len(table1) != len(table2):
            result.append(f"Table {t+1} row count differs: Original Doc has {len(table1)} rows vs Other Doc2 has {len(table2)} rows.")
    return result

# Streamlit app interface
st.title("Word Document Comparison Tool")
st.write("Upload two Word documents to compare texts and tables.")

file1 = st.file_uploader("Upload Original Document", type=["docx"])
file2 = st.file_uploader("Upload Other Document", type=["docx"])

if file1 and file2:
    st.write("Processing documents...")
    
    # Extract paragraphs (including content control text)
    paragraphs1 = extract_docx_paragraphs(file1)
    paragraphs2 = extract_docx_paragraphs(file2)
    
    # Extract tables (including content control text in cells)
    tables1 = extract_docx_tables(file1)
    tables2 = extract_docx_tables(file2)
    
    # Compare paragraphs and tables with detailed diff output and full changed values
    par_diff = compare_paragraphs(paragraphs1, paragraphs2)
    table_diff = compare_tables(tables1, tables2)
    
    # Create tabs for paragraph and table differences.
    tab1, tab2 = st.tabs(["Paragraph Differences", "Table Differences"])

    with tab1:
        if par_diff:
            for diff in par_diff:
                st.code(diff, language="diff")
        else:
            st.write("No differences in paragraphs detected.")
    
    with tab2:
        if table_diff:
            for diff in table_diff:
                st.code(diff, language="diff")
        else:
            st.write("No differences in tables detected.")
